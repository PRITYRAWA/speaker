'''
USING ONE DRIVE
get file and update merge fields and download that updated file in the form of pdf
'''
import datetime
import os
import subprocess
from os.path import join

import app_config
import identity.web
import pdfkit
import requests
from docx import Document
from docx2pdf import convert
from flask import Flask, jsonify, request, send_file, session
from flask_jwt_extended import (JWTManager, create_access_token,
                                create_refresh_token, jwt_required)
from flask_session import Session
from flask_sqlalchemy import SQLAlchemy
from passlib.hash import bcrypt

__version__ = "0.7.0"

app = Flask(__name__)
app.config.from_object(app_config)
assert app.config["REDIRECT_PATH"] != "/", "REDIRECT_PATH must not be /"
Session(app)
jwt = JWTManager(app) 

from werkzeug.middleware.proxy_fix import ProxyFix

app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

auth = identity.web.Auth(
    session=session,
    authority=app.config["AUTHORITY"],
    client_id=app.config["CLIENT_ID"],
    client_credential=app.config["CLIENT_SECRET"],
)

app.secret_key = "1GPkJM2AJOtRck5lJFDDlC1C2L-0VSbGmxKWx4sSSYY"
wkhtmltopdf_path = './wkhtmltopdf'
pdfkit_config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)

app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{join(app.root_path, 'project.db')}"
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(120), unique=True, nullable=False)
    
    def __init__(self, username, password):
        self.username = username
        self.password = bcrypt.hash(password)

    def verify_password(self, password):
        return bcrypt.verify(password, self.password)
    
with app.app_context():
    db.create_all()
    #hashed_password = bcrypt.hash("sharma95mansi")
    #print(hashed_password)
    # new_user = User(username='riyawalia@weboappdiscovery.onmicrosoft.com', password='onedrive@9814627072')

    # db.session.add(new_user)
    # db.session.commit()    

@app.route("/token", methods=["POST"])
def get_token():
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')

    user = User.query.filter_by(username=username).first()
    value=user.verify_password(password)
    # print("value",password)
    
    if user and user.verify_password(password):
        access_token = create_access_token(identity=user.id)
        refresh_token = create_refresh_token(identity=user.id)
        return jsonify(access_token=access_token, refresh_token=refresh_token), 200
    else:
        return jsonify(message="Invalid credentials"), 401
    
    
# SBC_SOW_LIVE_APPEARANCE_MAIN_(1) is represented as sslam1
@app.route("/sslam1", methods=["GET", "POST"])
@jwt_required()
def sslam1():
    
    user_inputs = request.json 
    
    client_id = app.config["CLIENT_ID"]
    client_secret = app.config["CLIENT_SECRET"]
    tenant_id = app.config["TENANT_ID"]
    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    token_data = {
        'grant_type': 'client_credentials',
        'client_id': client_id,
        'client_secret': client_secret,
        'scope': 'https://graph.microsoft.com/.default'
    }
    token_response = requests.post(token_url, data=token_data)
    print('token_response------------>',token_response)
    access_token = token_response.json().get('access_token')
    print('access_token---------->',access_token)
    
    user_id = "bb22a420-6705-4040-9064-2620019fbf38"
    
    drive_info_url = f'https://graph.microsoft.com/v1.0/users/{user_id}/drive'
    response = requests.get(drive_info_url, headers={'Authorization': f'Bearer {access_token}'})

    if response.status_code == 200:
        drive_info = response.json()
        drive_id = drive_info['id']
        print("Drive ID:", drive_id)
    else:
        print(f"Error: {response.status_code}")

    file_name = "SBC_SOW_LIVE_APPEARANCE_MAIN_(1).docx"  

    file_metadata_url = f'https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{file_name}'
    metadata_response = requests.get(file_metadata_url, headers={'Authorization': f'Bearer {access_token}'})

    if metadata_response.status_code == 200:
        file_metadata = metadata_response.json()
        file_size = file_metadata['size']
        file_download_url = file_metadata['@microsoft.graph.downloadUrl']

        file_response = requests.get(file_download_url)

        if file_response.status_code == 200:
            file_content = file_response.content

            local_file_path = f'{file_name}'
            with open(local_file_path, 'wb') as local_file:
                local_file.write(file_content)
            
            print(f"File '{file_name}' downloaded and saved as '{local_file_path}'")
            
    doc = Document(file_name)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for merge_field, replacement_value in user_inputs.items():
                if merge_field in text:
                    text = text.replace(merge_field, replacement_value)
            run.text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for merge_field, replacement_value in user_inputs.items():
                    if merge_field in text:
                        text = text.replace(merge_field, replacement_value)
                cell.text = text

    timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f") 
    
    new_file_name = f"{timestamp}_{file_name}"

    doc.save(new_file_name) 
        
    input_docx = new_file_name
    base_name = os.path.splitext(os.path.basename(input_docx))[0]
    output_pdf = f"{base_name}.pdf"
    
    # try:
    #     subprocess.run(['unoconv', '--output', output_pdf, '--format', 'pdf', input_docx], check=True)
    #     print(f"Conversion successful: {input_docx} -> {output_pdf}")
    # except subprocess.CalledProcessError:
    #     print(f"Conversion failed: {input_docx} -> {output_pdf}")

    # pdf_download_url = f"{request.url_root}download_pdf?filename={output_pdf}"
    
    # return {"pdf_download_url": pdf_download_url}
    
    try:
        subprocess.run(['/usr/bin/libreoffice', '--convert-to', 'pdf', input_docx])
        print(f"Conversion successful: {input_docx} -> {output_pdf}")

    except subprocess.CalledProcessError as e:
        print(f"Conversion failed: {input_docx} -> {output_pdf}, Error: {e}")
        return "fail"

    pdf_download_url = f"{request.url_root}download_pdf?filename={output_pdf}"
    return {"pdf_download_url": pdf_download_url}

@app.route("/sslam2", methods=["GET", "POST"])
@jwt_required()
def sslam2():
    
    user_inputs = request.json 
    
    client_id = app.config["CLIENT_ID"]
    client_secret = app.config["CLIENT_SECRET"]
    tenant_id = app.config["TENANT_ID"]
    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    token_data = {
        'grant_type': 'client_credentials',
        'client_id': client_id,
        'client_secret': client_secret,
        'scope': 'https://graph.microsoft.com/.default'
    }
    token_response = requests.post(token_url, data=token_data)
    print('token_response------------>',token_response)
    access_token = token_response.json().get('access_token')
    print('access_token---------->',access_token)
    
    user_id = "bb22a420-6705-4040-9064-2620019fbf38"
    
    drive_info_url = f'https://graph.microsoft.com/v1.0/users/{user_id}/drive'
    response = requests.get(drive_info_url, headers={'Authorization': f'Bearer {access_token}'})

    if response.status_code == 200:
        drive_info = response.json()
        drive_id = drive_info['id']
        print("Drive ID:", drive_id)
    else:
        print(f"Error: {response.status_code}")

    file_name = "SBC_SOW_LIVE_APPEARANCE_MAIN_(2).docx"  

    file_metadata_url = f'https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{file_name}'
    metadata_response = requests.get(file_metadata_url, headers={'Authorization': f'Bearer {access_token}'})

    if metadata_response.status_code == 200:
        file_metadata = metadata_response.json()
        file_size = file_metadata['size']
        file_download_url = file_metadata['@microsoft.graph.downloadUrl']

        file_response = requests.get(file_download_url)

        if file_response.status_code == 200:
            file_content = file_response.content

            local_file_path = f'{file_name}'
            with open(local_file_path, 'wb') as local_file:
                local_file.write(file_content)
            
            print(f"File '{file_name}' downloaded and saved as '{local_file_path}'")
            
    doc = Document(file_name)
    
    
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for merge_field, replacement_value in user_inputs.items():
                if merge_field in text:
                    text = text.replace(merge_field, replacement_value)
            run.text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for merge_field, replacement_value in user_inputs.items():
                    if merge_field in text:
                        text = text.replace(merge_field, replacement_value)
                cell.text = text

    timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f") 
    
    new_file_name = f"{timestamp}_{file_name}"

    doc.save(new_file_name) 
        
    input_docx = new_file_name
    base_name = os.path.splitext(os.path.basename(input_docx))[0]
    output_pdf = f"{base_name}.pdf"
    
    # try:
    #     subprocess.run(['unoconv', '--output', output_pdf, '--format', 'pdf', input_docx], check=True)
    #     print(f"Conversion successful: {input_docx} -> {output_pdf}")
    # except subprocess.CalledProcessError:
    #     print(f"Conversion failed: {input_docx} -> {output_pdf}")

    # pdf_download_url = f"{request.url_root}download_pdf?filename={output_pdf}"
    
    # return {"pdf_download_url": pdf_download_url}
    
    try:
        subprocess.run(['/usr/bin/libreoffice', '--convert-to', 'pdf', input_docx])
        print(f"Conversion successful: {input_docx} -> {output_pdf}")

    except subprocess.CalledProcessError as e:
        print(f"Conversion failed: {input_docx} -> {output_pdf}, Error: {e}")
        return "fail"

    pdf_download_url = f"{request.url_root}download_pdf?filename={output_pdf}"
    return {"pdf_download_url": pdf_download_url}

# SBC_SOW_GOC_LIVE_APPEARANCE_MAIN_(3) is represented as ssglam3
@app.route("/ssglam3", methods=["GET", "POST"])
@jwt_required()
async def ssglam3():
    
    user_inputs = request.json 
    
    client_id = app.config["CLIENT_ID"]
    client_secret = app.config["CLIENT_SECRET"]
    tenant_id = app.config["TENANT_ID"]
    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    token_data = {
        'grant_type': 'client_credentials',
        'client_id': client_id,
        'client_secret': client_secret,
        'scope': 'https://graph.microsoft.com/.default'
    }
    token_response = requests.post(token_url, data=token_data)
    access_token = token_response.json().get('access_token')
    
    user_id = "bb22a420-6705-4040-9064-2620019fbf38"
    
    drive_info_url = f'https://graph.microsoft.com/v1.0/users/{user_id}/drive'
    response = requests.get(drive_info_url, headers={'Authorization': f'Bearer {access_token}'})
    
    drive_id = None

    if response.status_code == 200:
        drive_info = response.json()
        drive_id = drive_info['id']
        print("Drive ID:", drive_id)
    else:
        print(f"Error: {response.status_code}")

    file_name = "SBC_SOW_GOC_LIVE_APPEARANCE_MAIN_(3).docx"  

    file_metadata_url = f'https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{file_name}'
    metadata_response = requests.get(file_metadata_url, headers={'Authorization': f'Bearer {access_token}'})

    if metadata_response.status_code == 200:
        file_metadata = metadata_response.json()
        file_size = file_metadata['size']
        file_download_url = file_metadata['@microsoft.graph.downloadUrl']

        file_response = requests.get(file_download_url)

        if file_response.status_code == 200:
            file_content = file_response.content

            local_file_path = f'{file_name}'
            with open(local_file_path, 'wb') as local_file:
                local_file.write(file_content)
            
            print(f"File '{file_name}' downloaded and saved as '{local_file_path}'")
    
    doc = Document(file_name)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for merge_field, replacement_value in user_inputs.items():
                if merge_field in text:
                    text = text.replace(merge_field, replacement_value)
            run.text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for merge_field, replacement_value in user_inputs.items():
                    if merge_field in text:
                        text = text.replace(merge_field, replacement_value)
                cell.text = text

    timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f") 
    
    new_file_name = f"{timestamp}_{file_name}"

    doc.save(new_file_name) 
        
    input_docx = new_file_name
    base_name = os.path.splitext(os.path.basename(input_docx))[0]
    output_pdf = f"{base_name}.pdf"

    # Convert the DOCX file to PDF using pyppeteer
    
    # try:
    #     subprocess.run(['unoconv', '--output', output_pdf, '--format', 'pdf', input_docx], check=True)
    #     print(f"Conversion successful: {input_docx} -> {output_pdf}")
    # except subprocess.CalledProcessError:
    #     print(f"Conversion failed: {input_docx} -> {output_pdf}")

    # pdf_download_url = f"{request.url_root}download_pdf?filename={output_pdf}"
    
    # return {"pdf_download_url": pdf_download_url}
    
    try:
        subprocess.run(['/usr/bin/libreoffice', '--convert-to', 'pdf', input_docx])
        print(f"Conversion successful: {input_docx} -> {output_pdf}")

    except subprocess.CalledProcessError as e:
        print(f"Conversion failed: {input_docx} -> {output_pdf}, Error: {e}")
        return "fail"

    pdf_download_url = f"{request.url_root}download_pdf?filename={output_pdf}"
    return {"pdf_download_url": pdf_download_url}




# SBC_SOW_VIRTUAL_MAIN_(4) is represented as ssvm4
@app.route("/ssvm4", methods=["GET", "POST"])
@jwt_required()
def ssvm4():
    
    user_inputs = request.json 
    
    client_id = app.config["CLIENT_ID"]
    client_secret = app.config["CLIENT_SECRET"]
    tenant_id = app.config["TENANT_ID"]
    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    token_data = {
        'grant_type': 'client_credentials',
        'client_id': client_id,
        'client_secret': client_secret,
        'scope': 'https://graph.microsoft.com/.default'
    }
    token_response = requests.post(token_url, data=token_data)
    access_token = token_response.json().get('access_token')
    
    user_id = "bb22a420-6705-4040-9064-2620019fbf38"
    
    drive_info_url = f'https://graph.microsoft.com/v1.0/users/{user_id}/drive'
    response = requests.get(drive_info_url, headers={'Authorization': f'Bearer {access_token}'})

    if response.status_code == 200:
        drive_info = response.json()
        drive_id = drive_info['id']
        print("Drive ID:", drive_id)
    else:
        print(f"Error: {response.status_code}")

    file_name = "SBC_SOW_VIRTUAL_MAIN_(4).docx"  

    file_metadata_url = f'https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{file_name}'
    metadata_response = requests.get(file_metadata_url, headers={'Authorization': f'Bearer {access_token}'})

    if metadata_response.status_code == 200:
        file_metadata = metadata_response.json()
        file_size = file_metadata['size']
        file_download_url = file_metadata['@microsoft.graph.downloadUrl']

        file_response = requests.get(file_download_url)

        if file_response.status_code == 200:
            file_content = file_response.content

            local_file_path = f'{file_name}'
            with open(local_file_path, 'wb') as local_file:
                local_file.write(file_content)
            
            print(f"File '{file_name}' downloaded and saved as '{local_file_path}'")
            
    doc = Document(file_name)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for merge_field, replacement_value in user_inputs.items():
                if merge_field in text:
                    text = text.replace(merge_field, replacement_value)
            run.text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for merge_field, replacement_value in user_inputs.items():
                    if merge_field in text:
                        text = text.replace(merge_field, replacement_value)
                cell.text = text

    timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f") 
    
    new_file_name = f"{timestamp}_{file_name}"

    doc.save(new_file_name) 
        
    input_docx = new_file_name
    base_name = os.path.splitext(os.path.basename(input_docx))[0]
    output_pdf = f"{base_name}.pdf"
    
    # try:
    #     subprocess.run(['unoconv', '--output', output_pdf, '--format', 'pdf', input_docx], check=True)
    #     print(f"Conversion successful: {input_docx} -> {output_pdf}")
    # except subprocess.CalledProcessError:
    #     print(f"Conversion failed: {input_docx} -> {output_pdf}")

    # pdf_download_url = f"{request.url_root}download_pdf?filename={output_pdf}"
    
    # return {"pdf_download_url": pdf_download_url}
    
    try:
        subprocess.run(['/usr/bin/libreoffice', '--convert-to', 'pdf', input_docx])
        print(f"Conversion successful: {input_docx} -> {output_pdf}")

    except subprocess.CalledProcessError as e:
        print(f"Conversion failed: {input_docx} -> {output_pdf}, Error: {e}")
        return "fail"

    pdf_download_url = f"{request.url_root}download_pdf?filename={output_pdf}"
    return {"pdf_download_url": pdf_download_url}






# SBC_SOW_GOC_VIRTUAL_APPEARANCE_MAIN_(5) is represented as ssgvam5
@app.route("/ssgvam5", methods=["GET", "POST"])
@jwt_required()
def ssgvam5():
    
    user_inputs = request.json 
    
    client_id = app.config["CLIENT_ID"]
    client_secret = app.config["CLIENT_SECRET"]
    tenant_id = app.config["TENANT_ID"]
    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    token_data = {
        'grant_type': 'client_credentials',
        'client_id': client_id,
        'client_secret': client_secret,
        'scope': 'https://graph.microsoft.com/.default'
    }
    token_response = requests.post(token_url, data=token_data)
    access_token = token_response.json().get('access_token')
    
    user_id = "bb22a420-6705-4040-9064-2620019fbf38"
    
    drive_info_url = f'https://graph.microsoft.com/v1.0/users/{user_id}/drive'
    response = requests.get(drive_info_url, headers={'Authorization': f'Bearer {access_token}'})

    if response.status_code == 200:
        drive_info = response.json()
        drive_id = drive_info['id']
        print("Drive ID:", drive_id)
    else:
        print(f"Error: {response.status_code}")

    file_name = "SBC_SOW_GOC_VIRTUAL_APPEARANCE_MAIN_(5).docx"  

    file_metadata_url = f'https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{file_name}'
    metadata_response = requests.get(file_metadata_url, headers={'Authorization': f'Bearer {access_token}'})

    if metadata_response.status_code == 200:
        file_metadata = metadata_response.json()
        file_size = file_metadata['size']
        file_download_url = file_metadata['@microsoft.graph.downloadUrl']

        file_response = requests.get(file_download_url)

        if file_response.status_code == 200:
            file_content = file_response.content

            local_file_path = f'{file_name}'
            with open(local_file_path, 'wb') as local_file:
                local_file.write(file_content)
            
            print(f"File '{file_name}' downloaded and saved as '{local_file_path}'")
            
    doc = Document(file_name)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for merge_field, replacement_value in user_inputs.items():
                if merge_field in text:
                    text = text.replace(merge_field, replacement_value)
            run.text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for merge_field, replacement_value in user_inputs.items():
                    if merge_field in text:
                        text = text.replace(merge_field, replacement_value)
                cell.text = text

    timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S%f") 
    
    new_file_name = f"{timestamp}_{file_name}"

    doc.save(new_file_name) 
        
    input_docx = new_file_name
    base_name = os.path.splitext(os.path.basename(input_docx))[0]
    output_pdf = f"{base_name}.pdf"
    
    # try:
    #     subprocess.run(['unoconv', '--output', output_pdf, '--format', 'pdf', input_docx], check=True)
    #     print(f"Conversion successful: {input_docx} -> {output_pdf}")
    # except subprocess.CalledProcessError:
    #     print(f"Conversion failed: {input_docx} -> {output_pdf}")

    # pdf_download_url = f"{request.url_root}download_pdf?filename={output_pdf}"
    
    # return {"pdf_download_url": pdf_download_url}
    
    try:
        subprocess.run(['/usr/bin/libreoffice', '--convert-to', 'pdf', input_docx])
        print(f"Conversion successful: {input_docx} -> {output_pdf}")

    except subprocess.CalledProcessError as e:
        print(f"Conversion failed: {input_docx} -> {output_pdf}, Error: {e}")
        return "fail"

    pdf_download_url = f"{request.url_root}download_pdf?filename={output_pdf}"
    return {"pdf_download_url": pdf_download_url}



@app.route("/download_pdf", methods=["GET"])
def download_pdf():
    filename = request.args.get("filename")
    pdf_path = f"{filename}"
    
    return send_file(pdf_path,
                     as_attachment=True,
                     download_name="output_file.pdf",
                     mimetype='application/pdf')

if __name__ == "__main__":
    app.run()



